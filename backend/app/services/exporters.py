"""Exportacao: TXT, DOCX, PDF, SRT, VTT, JSON, HTML, FCPXML.
Para Premiere Pro, o formato de importacao de legendas e o proprio SRT."""
import html
import io
import json


def _ts(seconds: float, sep: str = ",") -> str:
    h = int(seconds // 3600); m = int(seconds % 3600 // 60)
    s = int(seconds % 60); ms = int((seconds - int(seconds)) * 1000)
    return f"{h:02d}:{m:02d}:{s:02d}{sep}{ms:03d}"


def to_txt(segments: list) -> bytes:
    lines = []
    for s in segments:
        prefix = f"[{_ts(s['start'])[:8]}] "
        if s.get("speaker"):
            prefix += f"{s['speaker']}: "
        lines.append(prefix + s["text"])
    return "\n".join(lines).encode()


def to_srt(segments: list) -> bytes:
    out = []
    for i, s in enumerate(segments, 1):
        out.append(f"{i}\n{_ts(s['start'])} --> {_ts(s['end'])}\n{s['text']}\n")
    return "\n".join(out).encode()


def to_vtt(segments: list) -> bytes:
    out = ["WEBVTT\n"]
    for s in segments:
        out.append(f"{_ts(s['start'], '.')} --> {_ts(s['end'], '.')}\n{s['text']}\n")
    return "\n".join(out).encode()


def to_json(segments: list, meta: dict) -> bytes:
    return json.dumps({"meta": meta, "segments": segments}, ensure_ascii=False, indent=2).encode()


def to_html(segments: list, title: str) -> bytes:
    rows = "".join(
        f"<p><code>{_ts(s['start'])[:8]}</code> "
        + (f"<strong>{html.escape(s['speaker'])}:</strong> " if s.get("speaker") else "")
        + f"{html.escape(s['text'])}</p>"
        for s in segments
    )
    doc = (f"<!doctype html><html><head><meta charset='utf-8'><title>{html.escape(title)}</title>"
           "<style>body{font-family:Inter,sans-serif;max-width:720px;margin:40px auto;color:#111}"
           "code{font-family:'IBM Plex Mono',monospace;color:#6B6B6B}</style></head>"
           f"<body><h1>{html.escape(title)}</h1>{rows}</body></html>")
    return doc.encode()


def to_docx(segments: list, title: str) -> bytes:
    from docx import Document
    doc = Document()
    doc.add_heading(title, level=1)
    for s in segments:
        p = doc.add_paragraph()
        run = p.add_run(f"[{_ts(s['start'])[:8]}] ")
        run.font.name = "Courier New"
        if s.get("speaker"):
            p.add_run(f"{s['speaker']}: ").bold = True
        p.add_run(s["text"])
    buf = io.BytesIO(); doc.save(buf)
    return buf.getvalue()


def to_pdf(segments: list, title: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm)
    styles = getSampleStyleSheet()
    flow = [Paragraph(html.escape(title), styles["Title"]), Spacer(1, 12)]
    for s in segments:
        speaker = f"<b>{html.escape(s['speaker'])}:</b> " if s.get("speaker") else ""
        flow.append(Paragraph(
            f"<font face='Courier' color='#6B6B6B'>[{_ts(s['start'])[:8]}]</font> {speaker}{html.escape(s['text'])}",
            styles["BodyText"]))
    doc.build(flow)
    return buf.getvalue()


def to_fcpxml(segments: list, title: str, fps: int = 30) -> bytes:
    """FCPXML 1.9 com titles de legenda para Final Cut Pro."""
    def frames(sec): return f"{int(round(sec * fps))}/{fps}s"
    total = segments[-1]["end"] if segments else 1
    titles = "".join(
        f"<title ref='r2' offset='{frames(s['start'])}' duration='{frames(s['end']-s['start'])}' name='caption'>"
        f"<text><text-style ref='ts1'>{html.escape(s['text'])}</text-style></text></title>"
        for s in segments
    )
    xml = (
        "<?xml version='1.0' encoding='UTF-8'?><!DOCTYPE fcpxml><fcpxml version='1.9'>"
        f"<resources><format id='r1' frameDuration='1/{fps}s' width='1920' height='1080'/>"
        "<effect id='r2' name='Basic Title' uid='.../Titles.localized/Basic Title.moti'/></resources>"
        f"<library><event name='{html.escape(title)}'><project name='{html.escape(title)}'>"
        f"<sequence format='r1' duration='{frames(total)}'><spine>"
        f"<gap duration='{frames(total)}'>{titles}</gap>"
        "</spine></sequence></project></event></library></fcpxml>"
    )
    return xml.encode()


EXPORTERS = {
    "txt": ("text/plain", to_txt),
    "srt": ("application/x-subrip", to_srt),
    "vtt": ("text/vtt", to_vtt),
    "html": ("text/html", None),
    "json": ("application/json", None),
    "docx": ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", None),
    "pdf": ("application/pdf", None),
    "fcpxml": ("application/xml", None),
}
